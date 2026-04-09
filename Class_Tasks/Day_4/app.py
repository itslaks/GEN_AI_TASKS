import os
import json
import time
import uuid
import logging
import io
import pandas as pd
import pdfplumber
import docx
import httpx
import yaml
from dicttoxml import dicttoxml
from fpdf import FPDF
from flask import Flask, request, jsonify, send_from_directory, send_file
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger(__name__)

app = Flask(__name__, static_folder=".", static_url_path="")
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB limit

client = OpenAI(api_key=os.environ["OPENAI_API_KEY"])

SYSTEM_PROMPT = (
    "You are an expert data extraction agent. "
    "Your goal is to extract ALL relevant structured records from the provided text. "
    "Return ONLY a JSON object with a 'records' key containing an array of objects. "
    "Standard fields: {name, item, quantity, price, date (ISO-8601), currency (ISO-4217), confidence (0.0-1.0)}. "
    "If the user specifies 'Custom Fields', you MUST prioritize extracting those as well. "
    "Normalize: ₹/Rs/INR → INR, $→USD, €→EUR. "
    "Set missing fields to null. If multiple records/entities are found, extract ALL of them."
)


def _extract_mistral(text: str, custom_fields: str = None) -> dict:
    t0 = time.monotonic()
    
    prompt = SYSTEM_PROMPT
    if custom_fields:
        prompt += f"\n\nPriority Fields: {custom_fields}"
    
    payload = {
        "model": "mistral",
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ],
        "stream": False,
        "format": "json"
    }
    
    try:
        # Assuming Ollama is running on default port
        res = httpx.post("http://localhost:11434/api/chat", json=payload, timeout=60.0)
        res.raise_for_status()
        raw = res.json()["message"]["content"]
        data = json.loads(raw)
        latency_ms = round((time.monotonic() - t0) * 1000)
        return data, latency_ms, "mistral (local)", None
    except Exception as e:
        logger.error("mistral_fallback_failed error=%s", e)
        raise e


def _extract_json(text: str, custom_fields: str = None) -> dict:
    t0 = time.monotonic()
    
    user_content = text
    if custom_fields:
        user_content = f"Custom Fields to prioritize: {custom_fields}\n\nText to extract from:\n{text}"

    try:
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_content},
            ],
            max_tokens=2048,
            temperature=0,
            response_format={"type": "json_object"},
        )
        latency_ms = round((time.monotonic() - t0) * 1000)
        raw = response.choices[0].message.content
        data = json.loads(raw)
        usage = response.usage
        return data, latency_ms, response.model, usage
    except Exception as e:
        logger.warning("openai_failed_attempting_fallback error=%s", e)
        return _extract_mistral(text, custom_fields)


def _extract_text_from_file(file) -> str:
    filename = file.filename.lower()
    content = ""
    file_bytes = file.read()
    file_stream = io.BytesIO(file_bytes)

    try:
        if filename.endswith(".pdf"):
            with pdfplumber.open(file_stream) as pdf:
                content = "\n".join([page.extract_text() or "" for page in pdf.pages])
        elif filename.endswith(".docx"):
            doc = docx.Document(file_stream)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif filename.endswith((".xlsx", ".xls")):
            df = pd.read_excel(file_stream)
            content = df.to_string(index=False)
        elif filename.endswith(".csv"):
            df = pd.read_csv(file_stream)
            content = df.to_string(index=False)
        elif filename.endswith(".json"):
            data = json.load(file_stream)
            content = json.dumps(data, indent=2)
        else:  # Assume text
            content = file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        logger.error("file_parse_error filename=%s error=%s", filename, e)
        raise ValueError(f"Failed to parse {filename}: {str(e)}")

    return content.strip()


@app.route("/")
def index():
    return send_from_directory(".", "index.html")


@app.route("/upload", methods=["POST"])
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    try:
        text = _extract_text_from_file(file)
        return jsonify({"text": text, "filename": file.filename})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/export", methods=["POST"])
def export():
    body = request.get_json(silent=True) or {}
    data = body.get("data")
    fmt = body.get("format", "json").lower()

    if not data:
        return jsonify({"error": "Data is required"}), 400

    # Handle records list or single object
    records = []
    if isinstance(data, dict):
        if "records" in data:
            records = data["records"]
        else:
            records = [data]
    elif isinstance(data, list):
        records = data
    else:
        records = [data]

    try:
        output = io.BytesIO()
        if fmt == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(200, 10, txt="Extracted Data Report", ln=1, align='C')
            pdf.ln(10)
            
            if records:
                # Use all keys from the first record as headers
                headers = list(records[0].keys())
                pdf.set_font("Arial", 'B', 8)
                col_width = 190 / len(headers)
                
                for h in headers:
                    pdf.cell(col_width, 8, str(h).upper(), border=1)
                pdf.ln()
                
                pdf.set_font("Arial", '', 7)
                for rec in records:
                    for h in headers:
                        val = str(rec.get(h, ""))
                        pdf.cell(col_width, 7, val[:25], border=1) # Truncate long vals
                    pdf.ln()
            
            pdf_bytes = pdf.output()
            output.write(pdf_bytes)
            output.seek(0)
            return send_file(output, mimetype="application/pdf", as_attachment=True, download_name="data.pdf")

        elif fmt == "docx":
            doc = docx.Document()
            doc.add_heading('Extracted Data Report', 0)
            
            if records:
                headers = list(records[0].keys())
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Shading Accent 1'
                hdr_cells = table.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h).upper()
                
                for rec in records:
                    row_cells = table.add_row().cells
                    for i, h in enumerate(headers):
                        row_cells[i].text = str(rec.get(h, ""))
            
            doc.save(output)
            output.seek(0)
            return send_file(output, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document", as_attachment=True, download_name="data.docx")

        elif fmt == "xlsx":
            df = pd.DataFrame(records)
            df.to_excel(output, index=False)
            output.seek(0)
            return send_file(output, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", as_attachment=True, download_name="data.xlsx")

        elif fmt == "csv":
            df = pd.DataFrame(records)
            csv_str = df.to_csv(index=False)
            output.write(csv_str.encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="text/csv", as_attachment=True, download_name="data.csv")

        elif fmt == "tsv":
            df = pd.DataFrame(records)
            tsv_str = df.to_csv(index=False, sep='\t')
            output.write(tsv_str.encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="text/tab-separated-values", as_attachment=True, download_name="data.tsv")

        elif fmt == "json":
            json_str = json.dumps(data, indent=2)
            output.write(json_str.encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="application/json", as_attachment=True, download_name="data.json")

        elif fmt == "ndjson":
            lines = [json.dumps(r) for r in records]
            output.write("\n".join(lines).encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="application/x-ndjson", as_attachment=True, download_name="data.ndjson")

        elif fmt == "yaml":
            yaml_str = yaml.dump(data, sort_keys=False)
            output.write(yaml_str.encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="application/x-yaml", as_attachment=True, download_name="data.yaml")

        elif fmt == "xml":
            xml_bytes = dicttoxml(data, custom_root='nexus_extract', attr_type=False)
            output.write(xml_bytes)
            output.seek(0)
            return send_file(output, mimetype="application/xml", as_attachment=True, download_name="data.xml")

        elif fmt == "sql":
            if not records: return jsonify({"error": "No records"}), 400
            table_name = "extracted_records"
            cols = ", ".join(records[0].keys())
            sqls = []
            for r in records:
                vals = ", ".join([f"'{str(v).replace(chr(39), chr(39)+chr(39))}'" if v is not None else "NULL" for v in r.values()])
                sqls.append(f"INSERT INTO {table_name} ({cols}) VALUES ({vals});")
            output.write("\n".join(sqls).encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="text/x-sql", as_attachment=True, download_name="data.sql")

        elif fmt == "txt":
            lines = []
            for i, rec in enumerate(records):
                lines.append(f"--- Record {i+1} ---")
                for k, v in rec.items():
                    lines.append(f"{k.upper()}: {v}")
                lines.append("")
            output.write("\n".join(lines).encode('utf-8'))
            output.seek(0)
            return send_file(output, mimetype="text/plain", as_attachment=True, download_name="data.txt")

        else:
            return jsonify({"error": f"Unsupported format: {fmt}"}), 400

    except Exception as e:
        logger.error("export_error format=%s error=%s", fmt, e)
        return jsonify({"error": str(e)}), 500


@app.route("/extract", methods=["POST"])
def extract():
    request_id = str(uuid.uuid4())
    body = request.get_json(silent=True) or {}
    text = (body.get("text") or "").strip()
    custom_fields = body.get("customFields")

    if not text:
        logger.warning("request_id=%s empty input", request_id)
        resp = jsonify({"error": "text field is required", "request_id": request_id})
        resp.status_code = 400
        resp.headers["X-Request-Id"] = request_id
        return resp

    try:
        data, latency_ms, model, usage = _extract_json(text, custom_fields)

        total_tokens = usage.total_tokens if usage else None
        logger.info(
            "request_id=%s latency_ms=%d model=%s tokens=%s",
            request_id, latency_ms, model, total_tokens,
        )

        resp = jsonify({
            "data": data,
            "request_id": request_id,
            "latency_ms": latency_ms,
            "model": model,
            "tokens": total_tokens,
        })
        resp.headers["X-Request-Id"] = request_id
        return resp

    except json.JSONDecodeError as e:
        logger.error("request_id=%s json_parse_error=%s", request_id, e)
        resp = jsonify({"error": "Model returned invalid JSON", "request_id": request_id})
        resp.status_code = 502
        resp.headers["X-Request-Id"] = request_id
        return resp

    except Exception as e:
        logger.error("request_id=%s error=%s", request_id, e)
        resp = jsonify({"error": str(e), "request_id": request_id})
        resp.status_code = 500
        resp.headers["X-Request-Id"] = request_id
        return resp


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=os.environ.get("FLASK_DEBUG", "0") == "1")
